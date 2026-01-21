"""
文档信息提取工具

该工具可以分析PDF、Word和TXT文件并提取相关信息，
包括文件名、字符总数和页数（适用于PDF和Word文档）。
"""
import os
import sys
from pathlib import Path
from typing import List, Tuple, Dict, Optional

# 尝试导入所需的库
try:
    import PyPDF2
    from docx import Document
except ImportError as e:
    print(f"缺少必要的依赖库: {e}")
    print("请运行: pip install PyPDF2 python-docx")
    sys.exit(1)


def count_chars_in_pdf(file_path: str) -> int:
    """
    计算PDF文件中的字符数
    
    Args:
        file_path: PDF文件路径
        
    Returns:
        字符总数
    """
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            total_chars = 0
            
            for page in reader.pages:
                text = page.extract_text()
                total_chars += len(text)
                
        return total_chars
    except Exception as e:
        raise Exception(f"读取PDF文件时出错: {str(e)}")


def count_chars_and_pages_in_pdf(file_path: str) -> Tuple[int, int]:
    """
    计算PDF文件中的字符数和页数
    
    Args:
        file_path: PDF文件路径
        
    Returns:
        (字符总数, 页数)
    """
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            total_chars = 0
            total_pages = len(reader.pages)
            
            for page in reader.pages:
                text = page.extract_text()
                total_chars += len(text)
                
        return total_chars, total_pages
    except Exception as e:
        raise Exception(f"读取PDF文件时出错: {str(e)}")


def count_chars_and_pages_in_word(file_path: str) -> Tuple[int, int]:
    """
    计算Word文档中的字符数和页数
    
    Args:
        file_path: Word文档路径
        
    Returns:
        (字符总数, 页数估算)
    """
    try:
        doc = Document(file_path)
        total_chars = 0
        
        # 提取所有段落文本
        for paragraph in doc.paragraphs:
            total_chars += len(paragraph.text)
        
        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    total_chars += len(cell.text)
        
        # 估算页数（基于字符数，每页约2000字符）
        estimated_pages = max(1, (total_chars + 1999) // 2000)  # 向上取整
        
        return total_chars, estimated_pages
    except Exception as e:
        raise Exception(f"读取Word文档时出错: {str(e)}")


def count_chars_in_txt(file_path: str) -> int:
    """
    计算TXT文件中的字符数
    
    Args:
        file_path: TXT文件路径
        
    Returns:
        字符总数
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            return len(content)
    except UnicodeDecodeError:
        # 如果UTF-8解码失败，尝试其他编码
        try:
            with open(file_path, 'r', encoding='gbk') as file:
                content = file.read()
                return len(content)
        except UnicodeDecodeError:
            raise Exception(f"无法读取TXT文件，编码格式不支持: {file_path}")
    except Exception as e:
        raise Exception(f"读取TXT文件时出错: {str(e)}")


def get_file_info(file_path: str) -> Dict[str, any]:
    """
    获取单个文件的信息
    
    Args:
        file_path: 文件路径
        
    Returns:
        包含文件信息的字典
    """
    path_obj = Path(file_path)
    file_ext = path_obj.suffix.lower()
    
    info = {
        'file_name': path_obj.name,
        'char_count': 0,
        'page_count': '-'
    }
    
    if file_ext == '.pdf':
        char_count, page_count = count_chars_and_pages_in_pdf(file_path)
        info['char_count'] = char_count
        info['page_count'] = page_count
    elif file_ext == '.docx':
        char_count, page_count = count_chars_and_pages_in_word(file_path)
        info['char_count'] = char_count
        info['page_count'] = page_count
    elif file_ext == '.txt':
        char_count = count_chars_in_txt(file_path)
        info['char_count'] = char_count
    else:
        raise Exception(f"不支持的文件格式: {file_ext}")
    
    return info


def analyze_documents(file_paths: List[str]) -> List[Dict[str, any]]:
    """
    分析多个文档文件并提取信息
    
    Args:
        file_paths: 文件路径列表
        
    Returns:
        包含所有文件信息的列表
    """
    results = []
    
    for file_path in file_paths:
        try:
            if not os.path.exists(file_path):
                print(f"警告: 文件不存在 - {file_path}")
                continue
                
            info = get_file_info(file_path)
            results.append(info)
        except Exception as e:
            print(f"处理文件 {file_path} 时出错: {str(e)}")
            continue
    
    return results


def format_table(results: List[Dict[str, any]]) -> str:
    """
    将结果格式化为表格形式

    Args:
        results: 文件信息列表

    Returns:
        格式化的表格字符串
    """
    if not results:
        return "没有找到有效的文件或所有文件都无法处理"

    # 固定列宽以确保表格整齐
    name_width = 30
    char_width = 12
    page_width = 8

    # 构建表格
    table = []
    header = f"{'文件名':<{name_width}} │ {'字符数':>{char_width}} │ {'页数':>{page_width}}"
    separator = "─" * name_width + "─┼─" + "─" * char_width + "─┼─" + "─" * page_width

    table.append(f"┌{separator}┐")
    table.append(f"│ {header} │")
    table.append(f"├{separator}┤")

    for item in results:
        # 截断过长的文件名
        file_name = item['file_name'][:name_width-3] + "..." if len(item['file_name']) > name_width-3 else item['file_name']
        char_formatted = f"{item['char_count']:,}"  # 添加千位分隔符
        page_info = str(item['page_count'])

        row = f"│ {file_name:<{name_width}} │ {char_formatted:>{char_width}} │ {page_info:>{page_width}} │"
        table.append(row)

    table.append(f"└{separator}┘")

    return "\n".join(table)


def main():
    """
    主函数，演示如何使用文档分析工具
    """
    if len(sys.argv) < 2:
        print("使用方法: python document_analyzer.py <文件路径1> [文件路径2] ...")
        print("例如: python document_analyzer.py report.pdf proposal.docx notes.txt")
        return
    
    file_paths = sys.argv[1:]
    print("正在分析文档...")
    
    results = analyze_documents(file_paths)
    
    print("\n文档信息汇总:")
    print(format_table(results))


if __name__ == "__main__":
    main()