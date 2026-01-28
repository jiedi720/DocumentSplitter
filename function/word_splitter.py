"""
Word 分割逻辑

该模块实现了对 Word 文档（.docx）的各种分割方式，
包括按字符数分割和按段落数分割。
"""
import os
from docx import Document
from pathlib import Path
from .file_handler import FileHandler
from .chapter_detector import ChapterDetector


class WordSplitter:
    """Word 分割器

    该类提供了对 Word 文档（.docx）进行分割的功能，
    支持按字符数分割和按段落数分割两种方式。
    """

    def __init__(self):
        """初始化 Word 分割器

        创建文件处理器实例，用于处理通用文件操作。
        """
        self.file_handler = FileHandler()
        self.chapter_detector = ChapterDetector()

    def split_by_chars(self, input_path, chars_per_split, output_dir=None, preserve_chapter=False):
        """
        按字符数分割 Word 文件

        该方法将 Word 文档的内容提取为文本，然后按照指定的字符数进行分割，
        最后将每个文本部分保存为新的 Word 文档。

        Args:
            input_path (str): 输入 Word 文件的完整路径
            chars_per_split (int): 每个分割文件应包含的字符数
            output_dir (str, optional): 输出目录路径，默认为输入文件所在目录
            preserve_chapter (bool, optional): 是否保留章节完整性，默认为 False

        Returns:
            list: 包含所有成功分割的文件路径的列表

        Raises:
            FileNotFoundError: 当输入文件不存在时抛出
            ValueError: 当文件格式不正确或分割规则无效时抛出
        """
        # 如果未指定输出目录，则使用输入文件所在目录
        if output_dir is None:
            output_dir = str(Path(input_path).parent)

        # 验证输入文件是否存在
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"输入文件不存在: {input_path}")

        # 验证文件是否为有效的 DOCX 格式
        if not self.file_handler.get_file_type(input_path) == '.docx':
            raise ValueError(f"文件不是有效的 DOCX 格式: {input_path}")

        # 验证字符数分割规则是否有效
        if not self.file_handler.validate_split_rule(chars_per_split, '.docx'):
            raise ValueError(f"无效的字符数分割规则: {chars_per_split}")

        # 读取 Word 文档内容
        doc = Document(input_path)
        full_text = []  # 存储提取的文本内容

        # 提取所有段落文本
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)

        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        # 合并所有文本，用换行符分隔
        text_content = '\n'.join(full_text)

        # 检查是否需要分割：如果总字符数小于等于分割字符数，则复制整个文件
        if len(text_content) <= chars_per_split:
            output_path = self.file_handler.generate_output_filename(
                input_path, 1, '.docx'
            )
            new_doc = Document()
            for para in doc.paragraphs:
                new_doc.add_paragraph(para.text)
            new_doc.save(output_path)
            return [output_path]

        # 按字符数分割文本内容
        text_parts = []  # 存储分割后的文本片段

        if preserve_chapter:
            # 保留章节完整性：调整分割点
            current_pos = 0
            while current_pos < len(text_content):
                target_pos = min(current_pos + chars_per_split, len(text_content))

                # 调整分割点以避免在章节中间分割
                adjusted_pos = self.chapter_detector.adjust_split_point(
                    target_pos, text_content, 'cn', max_adjustment=chars_per_split // 2
                )

                # 确保至少分割一部分内容
                if adjusted_pos <= current_pos:
                    adjusted_pos = target_pos

                text_parts.append(text_content[current_pos:adjusted_pos])
                current_pos = adjusted_pos
        else:
            # 不保留章节完整性：直接按字符数分割
            for i in range(0, len(text_content), chars_per_split):
                text_parts.append(text_content[i:i + chars_per_split])

        # 将每个文本部分写入新的 Word 文件
        output_paths = []  # 存储输出文件路径的列表
        for idx, text_part in enumerate(text_parts, 1):
            # 生成输出文件名
            output_path = self.file_handler.generate_output_filename(
                input_path, idx, '.docx'
            )

            # 创建新的 Word 文档并写入内容
            new_doc = Document()

            # 尝试保持原始文档的样式
            # 这里简单地将文本按换行符分割成段落
            paragraphs = text_part.split('\n')
            for para_text in paragraphs:
                if para_text.strip():  # 避免添加只有空白字符的空段落
                    new_doc.add_paragraph(para_text)

            new_doc.save(output_path)
            output_paths.append(output_path)

        return output_paths

    def split_by_paragraphs(self, input_path, paras_per_split, output_dir=None, preserve_chapter=False):
        """
        按段落数分割 Word 文件

        该方法将 Word 文档按指定的段落数进行分割，
        每个分割后的文件包含指定数量的段落。

        Args:
            input_path (str): 输入 Word 文件的完整路径
            paras_per_split (int): 每个分割文件应包含的段落数
            output_dir (str, optional): 输出目录路径，默认为输入文件所在目录
            preserve_chapter (bool, optional): 是否保留章节完整性，默认为 False

        Returns:
            list: 包含所有成功分割的文件路径的列表

        Raises:
            FileNotFoundError: 当输入文件不存在时抛出
            ValueError: 当文件格式不正确或分割规则无效时抛出
        """
        # 如果未指定输出目录，则使用输入文件所在目录
        if output_dir is None:
            output_dir = str(Path(input_path).parent)

        # 验证输入文件是否存在
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"输入文件不存在: {input_path}")

        # 验证文件是否为有效的 DOCX 格式
        if not self.file_handler.get_file_type(input_path) == '.docx':
            raise ValueError(f"文件不是有效的 DOCX 格式: {input_path}")

        # 验证段落数分割规则是否有效
        if not self.file_handler.validate_split_rule(paras_per_split, '.docx'):
            raise ValueError(f"无效的段落数分割规则: {paras_per_split}")

        # 读取 Word 文档的所有段落
        doc = Document(input_path)
        all_paragraphs = doc.paragraphs

        # 检查是否需要分割：如果总段落数小于等于分割段落数，则复制整个文件
        if len(all_paragraphs) <= paras_per_split:
            output_path = self.file_handler.generate_output_filename(
                input_path, 1, '.docx'
            )
            new_doc = Document()
            for para in all_paragraphs:
                new_doc.add_paragraph(para.text)
            new_doc.save(output_path)
            return [output_path]

        # 如果需要保留章节完整性，先查找所有章节段落位置
        chapter_paras = []
        if preserve_chapter:
            chapter_paras = self.chapter_detector.find_paragraph_chapter_positions(all_paragraphs)
            chapter_para_indices = [ch['para_index'] for ch in chapter_paras]

        # 按段落数分割文档
        output_paths = []  # 存储输出文件路径的列表
        part_num = 1       # 分割部分的编号
        current_para = 0   # 当前处理到的段落索引

        while current_para < len(all_paragraphs):
            # 计算理论上的结束段落索引
            target_end_para = min(current_para + paras_per_split, len(all_paragraphs))

            # 如果需要保留章节完整性，调整结束段落索引
            if preserve_chapter and chapter_para_indices:
                # 查找目标结束段落之后的第一个章节段落
                next_chapter_para = None
                for chapter_para in chapter_para_indices:
                    if chapter_para > current_para and chapter_para < target_end_para:
                        next_chapter_para = chapter_para
                        break

                # 如果在目标范围内找到了章节段落，则在章节段落处分割
                if next_chapter_para is not None:
                    end_para = next_chapter_para
                else:
                    # 没有找到章节段落，使用原始的结束段落索引
                    end_para = target_end_para
            else:
                end_para = target_end_para

            # 生成输出文件名
            output_path = self.file_handler.generate_output_filename(
                input_path, part_num, '.docx'
            )

            # 创建新的 Word 文档并添加当前段落组
            new_doc = Document()

            # 复制原始文档的样式（简化版）
            for i in range(current_para, end_para):
                para = all_paragraphs[i]
                new_para = new_doc.add_paragraph(para.text)

                # 尝试复制段落的基本格式（简化版）
                if para.style:
                    new_para.style = para.style

            new_doc.save(output_path)
            output_paths.append(output_path)
            part_num += 1
            current_para = end_para

        return output_paths

    def split_by_equal_parts(self, input_path, parts_count, output_dir=None, preserve_chapter=False):
        """
        均分 Word 文件

        该方法将 Word 文档平均分割为指定的份数。

        Args:
            input_path (str): 输入 Word 文件的完整路径
            parts_count (int): 要分割的份数
            output_dir (str, optional): 输出目录路径，默认为输入文件所在目录
            preserve_chapter (bool, optional): 是否保留章节完整性，默认为 False

        Returns:
            list: 包含所有成功分割的文件路径的列表

        Raises:
            FileNotFoundError: 当输入文件不存在时抛出
            ValueError: 当文件格式不正确或分割规则无效时抛出
        """
        # 如果未指定输出目录，则使用输入文件所在目录
        if output_dir is None:
            output_dir = str(Path(input_path).parent)

        # 验证输入文件是否存在
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"输入文件不存在: {input_path}")

        # 验证文件是否为有效的 DOCX 格式
        if not self.file_handler.get_file_type(input_path) == '.docx':
            raise ValueError(f"文件不是有效的 DOCX 格式: {input_path}")

        # 验证分割份数是否有效
        if parts_count <= 0:
            raise ValueError("分割份数必须大于 0")

        # 读取 Word 文档内容
        doc = Document(input_path)
        full_text = []  # 存储提取的文本内容

        # 提取所有段落文本
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)

        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        # 合并所有文本，用换行符分隔
        text_content = '\n'.join(full_text)
        total_chars = len(text_content)

        # 计算每份的字符数
        if parts_count >= total_chars:
            # 如果份数大于等于总字符数，每个文件至少一个字符
            chars_per_part = 1
        else:
            # 计算每份的字符数，向上取整
            chars_per_part = (total_chars + parts_count - 1) // parts_count

        # 调用 split_by_chars 方法进行分割
        return self.split_by_chars(input_path, chars_per_part, output_dir, preserve_chapter)

    def merge_docs(self, input_files, output_path=None):
        """
        合并多个 Word 文件

        该方法将多个 Word 文档合并为一个单一的 Word 文档。

        Args:
            input_files (list): 要合并的 Word 文件路径列表
            output_path (str, optional): 输出文件路径，默认为自动生成

        Returns:
            str: 合并后文件的路径

        Raises:
            FileNotFoundError: 当输入文件不存在时抛出
            ValueError: 当文件格式不正确或输入列表为空时抛出
        """
        # 验证输入文件列表
        if not input_files:
            raise ValueError("输入文件列表不能为空")

        # 验证所有输入文件是否为 DOCX 格式
        for file_path in input_files:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"输入文件不存在: {file_path}")
            if not self.file_handler.get_file_type(file_path) == '.docx':
                raise ValueError(f"文件不是有效的 DOCX 格式: {file_path}")

        # 如果未指定输出路径，则自动生成
        if output_path is None:
            output_path = self.file_handler.generate_merge_output_filename(input_files)

        # 创建新的 Word 文档作为合并结果
        merged_doc = Document()

        # 合并所有 Word 文件
        for file_path in input_files:
            doc = Document(file_path)
            
            # 复制所有段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    new_para = merged_doc.add_paragraph(paragraph.text)
                    # 尝试复制段落样式
                    if paragraph.style:
                        new_para.style = paragraph.style
            
            # 复制所有表格
            for table in doc.tables:
                new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                # 复制表格内容
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.cell(i, j).text = cell.text

        # 保存合并后的文档
        merged_doc.save(output_path)

        return output_path